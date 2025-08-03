from const import (
    column_mappings,
    convert_str_field,
    excel_file_name,
    excel_sheet_name,
    word_file_name,
)
import os
import re
import argparse
import pandas as pd
from docx import Document
from utils import replace_placeholder_preserve_style

output_dir = "generated_docs_2"
os.makedirs(output_dir, exist_ok=True)

parser = argparse.ArgumentParser()
parser.add_argument("--group-by", type=str, required=True)
args = parser.parse_args()

# Read Excel file (first sheet by default)
df = pd.read_excel(
    excel_file_name,
    sheet_name=excel_sheet_name,
    engine="openpyxl",
    converters={column_mappings[field]: str for field in convert_str_field},
)

result = {}

# For each row in the Excel file
for index, row in df.iterrows():
    # Load template

    temp = {}

    for var, col in column_mappings.items():
        if col in row and pd.notna(row[col]):
            value = row[col]
            # Check if it's a number (int or float), format with digit grouping
            if isinstance(value, (int, float)) and "tiền" in col:
                formatted_value = f"{value:,.0f}".replace(",", ".")
            else:
                formatted_value = str(value)

            temp[f"{{{{{var}}}}}"] = formatted_value

    result.setdefault(row[column_mappings[args.group_by]], []).append(temp)

for group_key, value in result.items():
    doc = Document(word_file_name)

    print("Key:", group_key)
    print("Value:", len(value))

    replacements = value[0]

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                replace_placeholder_preserve_style(paragraph, key, value)

    for table in doc.tables:
        for row_obj in table.rows:
            for cell in row_obj.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, value)

    # Save with a unique name (e.g., by borrower name or row number)
    group_key = str(group_key).replace(" ", "_").replace("/", "_")

    filename = f"{group_key}_2.docx"
    doc.save(os.path.join(output_dir, filename))

print("All documents generated successfully.")
