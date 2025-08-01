import pandas as pd
from docx import Document
import os
from datetime import datetime
import re
from num2words import num2words
from dateutil.relativedelta import relativedelta
from const import (
    excel_file_name,
    excel_sheet_name,
    column_mappings,
    convert_str_field,
    word_file_name,
)


# Output folder
output_dir = "generated_docs"
os.makedirs(output_dir, exist_ok=True)

today = datetime.today()
current_date = f"ngày {today.day:02d} tháng {today.month:02d} năm {today.year}"

# Read Excel file (first sheet by default)
df = pd.read_excel(
    excel_file_name,
    sheet_name=excel_sheet_name,
    engine="openpyxl",
    converters={column_mappings[field]: str for field in convert_str_field},
)

print(df.dtypes)


def evaluate_placeholder(expr: str):
    match = re.match(r"{{\s*(\w+)\((\w+)\)\s*}}", expr)
    if not match:
        return None

    func_name, arg_name = match.groups()

    return func_name, arg_name


def generate_ky_han_tra_no_goc(start_date, total_duration_months, amount, cycle_months):
    start_date = datetime.strptime(start_date, "%d/%m/%Y")

    # Date list
    date_list = []
    elapsed_months = 0
    while elapsed_months < total_duration_months - cycle_months:
        new_date = start_date + relativedelta(months=elapsed_months)
        date_list.append(new_date)
        if elapsed_months + cycle_months >= total_duration_months - cycle_months:
            new_date = new_date + relativedelta(
                months=(total_duration_months - elapsed_months)
            )
            date_list.append(new_date)
        elapsed_months += cycle_months

    # Money list
    cycle_count = total_duration_months // cycle_months
    money_each_cycle = amount * cycle_months / total_duration_months
    money_list = [money_each_cycle] * cycle_count
    spare_money = amount - (money_each_cycle * cycle_count)
    if spare_money > 0:
        money_list.append(spare_money)

    # Combine two lists to get result
    result_lines = []

    for index, date in enumerate(date_list):
        amount_str = f"{(money_list[index]):,.0f}".replace(".", ",")
        result_lines.append(
            f"\t- Ngày {date.strftime('%d/%m/%Y')}, số tiền {amount_str} đồng."
        )

    # Combine all into one string
    final_output = "\n".join(result_lines)

    return final_output


# For each row in the Excel file
for index, row in df.iterrows():
    # Load template
    doc = Document(word_file_name)

    replacements = {}

    for var, col in column_mappings.items():
        if col in row and pd.notna(row[col]):
            value = row[col]

            # Check if it's a number (int or float), format with digit grouping
            if isinstance(value, (int, float)) and "tiền" in col:
                formatted_value = f"{value:,.0f}".replace(".", ",")
            else:
                formatted_value = str(value)

            replacements[f"{{{{{var}}}}}"] = formatted_value

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        if f"{{ky_han_tra_no_goc}}" in paragraph.text:
            p = paragraph._element

            paragraph.insert_paragraph_before(
                generate_ky_han_tra_no_goc(
                    amount=row[column_mappings["so_tien_cho_vay_dong"]],
                    start_date=row[column_mappings["ngay_bat_dau_tra_goc"]],
                    cycle_months=row[column_mappings["ky_han_tra_no_goc_so_thang_ky"]],
                    total_duration_months=row[
                        column_mappings["thoi_gian_cho_vay_thang"]
                    ],
                )
                # generate_ky_han_tra_no_goc(
                #     amount=100,
                #     cycle_months=12,
                #     start_date="22/05/2025",
                #     total_duration_months=50,
                # )
            )

            p.getparent().remove(p)

        for run in paragraph.runs:
            if f"{{current_date}}" in run.text:
                run.text = run.text.replace("{{current_date}}", str(current_date))

            # Replace function placeholders
            func_placeholders = re.findall(r"{{\s*\w+\(\w+\)\s*}}", run.text)
            for ph in func_placeholders:
                func_name, arg_name = evaluate_placeholder(ph)
                if func_name == "num2words":
                    run.text = run.text.replace(
                        ph,
                        num2words(
                            row[column_mappings[arg_name]], lang="vi"
                        ).capitalize()
                        + " đồng",
                    )

            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row_obj in table.rows:
            for cell in row_obj.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, value)

    # Save with a unique name (e.g., by borrower name or row number)
    borrower = (
        str(row[column_mappings["ho_ten_nguoi_vay"]])
        .replace(" ", "_")
        .replace("/", "_")
    )

    filename = f"{borrower}.docx"
    doc.save(os.path.join(output_dir, filename))

print("All documents generated successfully.")
